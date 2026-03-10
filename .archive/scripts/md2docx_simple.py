#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Simple Markdown to Word converter"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def convert_md_to_docx(md_file, docx_file):
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    doc = Document()
    
    # Set Chinese font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')
        stripped = line.strip()
        
        if not stripped:
            i += 1
            continue
        
        # H1
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(24)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # H2
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(18)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # H3
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(14)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # H4
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Blockquote
        elif stripped.startswith('> '):
            p = doc.add_paragraph(stripped[2:])
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run.font.italic = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Code block - skip
        elif stripped.startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
        
        # Table
        elif '|' in stripped and stripped.count('|') >= 3:
            table_lines = []
            start = i
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            i -= 1
            
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.split('|')]
                if cells and cells[0] == '': cells = cells[1:]
                if cells and cells[-1] == '': cells = cells[:-1]
                if cells and not all(c.startswith('-') for c in cells):
                    rows.append(cells)
            
            if len(rows) >= 2:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Table Grid'
                for ri, row_data in enumerate(rows):
                    for ci, cell_text in enumerate(row_data):
                        if ci < len(table.rows[ri].cells):
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    run.font.name = 'Microsoft YaHei'
                                    run.font.size = Pt(10)
                                    if ri == 0:
                                        run.font.bold = True
                                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Unordered list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            p = doc.add_paragraph(stripped[2:], style='List Bullet')
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Ordered list
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(text, style='List Number')
            for run in p.runs:
                run.font.name = 'Microsoft YaHei'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Horizontal rule
        elif stripped.startswith('---') or stripped.startswith('***'):
            p = doc.add_paragraph('_' * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Regular paragraph
        else:
            # Handle bold
            if '**' in stripped:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                p = doc.add_paragraph(stripped)
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        i += 1
    
    doc.save(docx_file)
    print(f"✓ 转换成功：{docx_file}")

if __name__ == '__main__':
    md_file = '/Users/jo/.openclaw/workspace/2026 年度战略规划 - 王健豪.md'
    docx_file = '/Users/jo/.openclaw/workspace/2026 年度战略规划 - 王健豪.docx'
    convert_md_to_docx(md_file, docx_file)
