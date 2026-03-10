#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convert Markdown to Word (.docx) with proper Chinese support
"""

import markdown
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys

def set_cell_background(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def convert_md_to_docx(md_file, docx_file):
    """Convert Markdown file to Word document"""
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set default font for Chinese support
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Parse markdown to HTML for processing
    md = markdown.Markdown(extensions=['tables', 'fenced_code', 'codehilite'])
    html = md.convert(content)
    
    # Process line by line
    lines = content.split('\n')
    current_list = []
    in_list = False
    list_type = None
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty lines
        if not stripped:
            i += 1
            continue
        
        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=1)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(24)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=2)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(18)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=3)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(14)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        elif stripped.startswith('#### '):
            p = doc.add_heading(stripped[5:], level=4)
            run = p.runs[0] if p.runs else p.add_run()
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(12)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Blockquote
        elif stripped.startswith('> '):
            quote_text = stripped[2:]
            # Check if it's a special quote with bold
            if quote_text.startswith('**') and '**' in quote_text[2:]:
                # Extract bold text
                p = doc.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                # Add italic formatting for blockquote
                run = p.add_run(quote_text)
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(11)
                run.font.italic = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                p = doc.add_paragraph()
                run = p.add_run(quote_text)
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(11)
                run.font.italic = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Code blocks (fenced)
        elif stripped.startswith('```'):
            # Skip the code block, collect lines until closing ```
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                # Set background by using shading (simplified)
                p.paragraph_format.space_after = Pt(6)
            continue
        
        # Tables
        elif '|' in stripped and stripped.count('|') >= 3:
            # This is a table, collect all table lines
            table_lines = []
            table_start = i
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Will be incremented at end of loop
            
            if len(table_lines) >= 2:
                # Parse table
                rows = []
                for tline in table_lines:
                    cells = [c.strip() for c in tline.split('|')]
                    # Remove first and last if empty
                    if cells and cells[0] == '':
                        cells = cells[1:]
                    if cells and cells[-1] == '':
                        cells = cells[:-1]
                    if cells and not all(c.startswith('-') for c in cells):
                        rows.append(cells)
                
                if len(rows) >= 2:
                    # Create table
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        row = table.rows[row_idx]
                        for col_idx, cell_text in enumerate(row_data):
                            if col_idx < len(row.cells):
                                cell = row.cells[col_idx]
                                cell.text = cell_text
                                # Bold first row (header)
                                if row_idx == 0:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.name = 'Microsoft YaHei'
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
                                else:
                                    for paragraph in cell.paragraphs:
                                        for run in paragraph.runs:
                                            run.font.name = 'Microsoft YaHei'
                                            run.font.size = Pt(10)
                                            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            continue
        
        # Unordered list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            list_text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(list_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Ordered list (numbered)
        elif re.match(r'^\d+\.\s', stripped):
            list_text = re.sub(r'^\d+\.\s', '', stripped)
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(list_text)
            run.font.name = 'Microsoft YaHei'
            run.font.size = Pt(11)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Separator line
        elif stripped.startswith('---') or stripped.startswith('***'):
            p = doc.add_paragraph()
            run = p.add_run('_' * 50)
            run.font.color.rgb = None  # Use default color
        
        # Regular paragraph
        else:
            # Check for bold text
            if '**' in stripped:
                p = doc.add_paragraph()
                parts = re.split(r'(\*\*.*?\*\*)', stripped)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.font.bold = True
                    else:
                        run = p.add_run(part)
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(11)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            else:
                p = doc.add_paragraph(stripped)
                run = p.runs[0] if p.runs else p.add_run()
                run.font.name = 'Microsoft YaHei'
                run.font.size = Pt(11)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        i += 1
    
    # Save the document
    doc.save(docx_file)
    print(f"✓ 成功转换：{docx_file}")

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        convert_md_to_docx(sys.argv[1], sys.argv[2])
    else:
        # Default paths
        md_file = '/Users/jo/.openclaw/workspace/2026 年度战略规划 - 王健豪.md'
        docx_file = '/Users/jo/.openclaw/workspace/2026 年度战略规划 - 王健豪.docx'
        convert_md_to_docx(md_file, docx_file)
